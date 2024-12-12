
## imgae file 

from google.oauth2 import service_account
from googleapiclient.discovery import build
from apiclient.http import MediaFileUpload, MediaIoBaseDownload
from deep_translator import GoogleTranslator
from langdetect import detect
import io
import concurrent.futures
import os
import time
import json
import hashlib
import pdfplumber
import pandas as pd
from docx import Document
from tabulate import tabulate
from multiprocessing import Pool
import xlrd
from openpyxl import load_workbook
import re
import logging
import random
import string
import subprocess
from bs4 import BeautifulSoup
# from docx import Document
from datetime import datetime,timedelta
from pdf2image import convert_from_path

# Configure logging
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger(__name__)

# Function to calculate the hash of a file
def calculate_file_hash(file_path):
    with open(file_path, "rb") as f:
        file_hash = hashlib.sha256(f.read()).hexdigest()
    return file_hash

# Function to remove non-ASCII characters from a string
def remove_non_ascii(text):
    return ''.join(i for i in text if ord(i) < 128 or ord(i) > 127)

# Function to extract data from PDF
def extract_from_pdf(pdf_path):
    try:
        all_text_data = ""
        all_table_data = ""

        with pdfplumber.open(pdf_path) as pdf:
            for page in pdf.pages:
                # Extract text
                text_data = page.extract_text()
                all_text_data += f"{text_data}\n\n"

                # Extract tables
                tables = page.extract_tables()
                if tables:
                    for table in tables:
                        df = pd.DataFrame(table[1:], columns=table[0])
                        all_table_data += tabulate(df, headers='keys', tablefmt='grid') + "\n\n"

        return remove_non_ascii(all_text_data), remove_non_ascii(all_table_data)
    except Exception as e:
        return f"Error extracting PDF: {e}", ""

# Modified process_document function to include table extraction for DOCX files
def process_document(file_path):
    def convert_doc_to_txt(file_path, output_file):
        # Convert .doc to .txt using LibreOffice
        subprocess.run(['libreoffice', '--headless', '--convert-to', 'txt:Text', '--outdir', os.path.dirname(output_file), file_path])
        os.rename(file_path.replace('.doc', '.txt'), output_file)

    def read_docx(file_path1):
        # Read text and tables from .docx file
        doc = Document(file_path1)
        text = '\n'.join([paragraph.text for paragraph in doc.paragraphs])

        # Extract tables
        tables_data = []
        for table in doc.tables:
            table_data = []
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells]
                table_data.append(row_data)
            tables_data.append(table_data)
        
        tables_text = "\n\n".join([tabulate(table, tablefmt="grid") for table in tables_data])
        return text, tables_text

    def read_txt(file_path1):
        # Read text from .txt file
        with open(file_path1, 'r', encoding='utf-8') as file:
            return file.read()

    output_file = file_path.replace('.doc', '.txt')

    if file_path.endswith('.docx'):
        text, tables_text = read_docx(file_path)
        return text, tables_text
    elif file_path.endswith('.doc'):
        convert_doc_to_txt(file_path, output_file)
        text = read_txt(output_file)
        # os.remove(output_file)
        return text, ""
    else:
        raise ValueError("Unsupported file format. Only .doc and .docx are supported.")
    
## function to extract data from image file

def googleOcr(imgfile):
    try:
        creds = service_account.Credentials.from_service_account_file(
            r"/data/imageExtraction/GoogleAPICred/projectoct-436907-a6e51afb9d49.json",
            scopes=['https://www.googleapis.com/auth/drive']
        )
        service = build('drive', 'v3', credentials=creds)
        mime = 'application/vnd.google-apps.document'
        res = service.files().create(
            body={
                'name': imgfile,
                'mimeType': mime
            },
            media_body=MediaFileUpload(imgfile, mimetype=mime, resumable=True)
        ).execute()
        
        text_output = io.BytesIO()
        downloader = MediaIoBaseDownload(
            text_output,
            service.files().export_media(fileId=res['id'], mimeType="text/plain")
        )
        done = False
        while not done:
            status, done = downloader.next_chunk()
        text_output.seek(0)
        extracted_text = text_output.read().decode('utf-8')
        service.files().delete(fileId=res['id']).execute()
        
        return extracted_text

    except Exception as e:
        logger.info(e)
        return None

def process_image(file_path: str) -> str:
    """
    Process the uploaded image file to extract text using OCR.
    """
    extracted_text = googleOcr(file_path)
    return extracted_text if extracted_text else "No text extracted from image."

## image pagefile 


def googleOcr1(file_path):
    try:
        creds = service_account.Credentials.from_service_account_file(
            r"/data/imageExtraction/GoogleAPICred/projectoct-436907-a6e51afb9d49.json",
            scopes=['https://www.googleapis.com/auth/drive']
        )
        service = build('drive', 'v3', credentials=creds)
        
        # Handle PDF files by converting them to images first
        if file_path.lower().endswith('.pdf'):
            images = convert_from_path(file_path)
            extracted_text = ""
            filename = os.path.basename(file_path).split('.')[0]
            for i, img in enumerate(images):
                img_path = f"{filename}_page_{i}.png"
                img.save(img_path, 'PNG')  # Save each page as an image
                page_text = googleOcr(img_path)  # Perform OCR on each page
                extracted_text += (page_text or "") + "\n"  # Handle None return case
                full_img_path = os.path.join(os.getcwd(), img_path)
                os.remove(full_img_path)
            return extracted_text.strip()

        mime = 'application/vnd.google-apps.document'
        res = service.files().create(
            body={'name': os.path.basename(file_path), 'mimeType': mime},
            media_body=MediaFileUpload(file_path, mimetype=mime, resumable=True)
        ).execute()

        text_output = io.BytesIO()
        downloader = MediaIoBaseDownload(text_output, service.files().export_media(fileId=res['id'], mimeType="text/plain"))
        
        done = False
        while not done:
            status, done = downloader.next_chunk()

        text_output.seek(0)
        extracted_text = text_output.read().decode('utf-8')
        service.files().delete(fileId=res['id']).execute()
        
        return extracted_text

    except Exception as e:
        logger.info(f"Error in OCR: {e}")
        return None




# Function to extract data from CSV
def extract_from_csv(csv_path):
    try:
        data = pd.read_csv(csv_path, encoding='utf-8')
        return remove_non_ascii(data.to_string()), ""
    except FileNotFoundError as e:
        return f"Error extracting CSV: FileNotFoundError: {e}", ""
    except pd.errors.EmptyDataError as e:
        return f"Error extracting CSV: EmptyDataError: {e}", ""
    except Exception as e:
        return f"Error extracting CSV: {e}", ""

def extract_from_excel(file_path):
    try:
        data = ""

        if file_path.lower().endswith(".xls"):
            wb = xlrd.open_workbook(file_path)
            sheet = wb.sheet_by_index(0)
            for row in range(sheet.nrows):
                row_data = sheet.row_values(row)
                # Clean trailing empty cells at the end of each row
                while row_data and not row_data[-1]:
                    row_data.pop()
                data += " | ".join(map(str, row_data)) + "\n"
        
        elif file_path.lower().endswith(".xlsx"):
            wb = load_workbook(file_path)
            sheet = wb.active
            for row in sheet.iter_rows(values_only=True):
                # Clean trailing empty cells at the end of each row
                row_data = list(row)
                while row_data and not row_data[-1]:
                    row_data.pop()
                data += " | ".join(map(str, row_data)) + "\n"

        # Return the raw data including non-ASCII characters
        return data, ""
    except Exception as e:
        return f"Error extracting Excel: {e}", ""

# Function to extract data from HTML
def extract_from_html(html_path):
    try:
        with open(html_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            text = soup.get_text(separator='\n')
            return remove_non_ascii(text), ""
    except Exception as e:
        return f"Error extracting HTML: {e}", ""

# Function to process a single file
def process_file(file_path, folder_data):
    filename = os.path.basename(file_path)
    file_hash = calculate_file_hash(file_path)

    file_data = {"text": "", "tables": "", "error": ""}

    try:
        if filename.lower().endswith(".pdf"):
            text_data, table_data = extract_from_pdf(file_path)
            file_data["text"] = text_data
            file_data["tables"] = table_data
            # extracted_text =googleOcr1(file_path)
            # file_data["text"] = extracted_text
        elif filename.lower().endswith(".docx") or filename.lower().endswith(".doc"):
            text_data, table_data = process_document(file_path)
            file_data["text"] = text_data
            file_data["tables"] = table_data
        # elif filename.lower().endswith(".csv"):
            # text_data, _ = extract_from_csv(file_path)
            # file_data["text"] = text_data
        # elif filename.lower().endswith(".xls") or filename.lower().endswith(".xlsx"):
            # text_data, _ = extract_from_excel(file_path)
            # file_data["text"] = text_data
        elif filename.lower().endswith((".jpg", ".jpeg", ".png",".gif")):
                extracted_text = process_image(file_path)
                file_data["text"] = extracted_text
            
        elif filename.lower().endswith(".html"):
            text_data, _ = extract_from_html(file_path)
            file_data["text"] = text_data
        else:
            file_data["error"] = "Unsupported file format"
    except FileNotFoundError:
        file_data["error"] = f"File not found: {file_path}"
    except Exception as e:
        file_data["error"] = f"Error processing {filename}: {e}"

    if file_hash in folder_data:
        new_filename = f"{filename}_{''.join(random.choices(string.ascii_letters + string.digits, k=6))}"
        folder_data[new_filename] = file_data
    else:
        folder_data[filename] = file_data

    # if not file_data["error"]:
    #     os.remove(file_path)

def process_directory(directory_path, output_base_path, timings_file):
    folder_timings = {}

    # Map to track merged folders by their 8-digit prefix
    merged_folders = {}

    for root, dirs, files in os.walk(directory_path):
        folder_start_time = time.time()
        folder_data = {}
        folder_count = 0

        # Extract the first 8 digits from the directory name
        match = re.match(r"(\d{8})", os.path.basename(root))
        if not match:
            continue

        folder_prefix = match.group(1)
        
        if folder_prefix not in merged_folders:
            final_dir_name = folder_prefix
            output_dir = os.path.join(output_base_path, final_dir_name)
            os.makedirs(output_dir, exist_ok=True)
            output_file = os.path.join(output_dir, f'{final_dir_name}.txt')
            merged_folders[folder_prefix] = output_file
        else:
            output_file = merged_folders[folder_prefix]

        if os.path.exists(output_file):
            print(f"Appending to {output_file}.")
        else:
            print(f"Creating {output_file}.")

        for filename in files:
            file_path = os.path.join(root, filename)
            process_file(file_path, folder_data)
            folder_count += 1

        try:
            with open(output_file, 'a', encoding='utf-8') as f:
                for filename, file_data in folder_data.items():
                    f.write(f"File: {filename}\n\n")
                    if file_data["text"]:
                        f.write("Text Data:\n")
                        f.write(file_data["text"] + "\n\n")
                    if file_data["tables"]:
                        f.write("Table Data:\n")
                        f.write(file_data["tables"] + "\n\n")
                    if file_data["error"]:
                        f.write("Error:\n")
                        f.write(file_data["error"] + "\n\n")
                    f.write("-" * 50 + "\n\n")
        except FileNotFoundError as e:
            print(f"Error: {e}")
            continue

        folder_elapsed_time = time.time() - folder_start_time
        folder_timings[root] = folder_elapsed_time

        with open(timings_file, 'a') as f:
            f.write(f"{root}: Processed {folder_count} files in {folder_elapsed_time:.2f} seconds.\n")

    #     # Always attempt to remove the directory after processing
    #     try:
    #         os.rmdir(root)
    #     except OSError as e:
    #         print(f"Error removing directory {root}: {e}")

    # return folder_timings

def save_processed_files(processed_files, processed_files_path):
    with open(processed_files_path, 'w') as f:
        json.dump(list(processed_files), f)

# def load_processed_files(processed_files_path):
#     if os.path.exists(processed_files_path):
#         with open(processed_files_path, 'r') as f:
#             return set(json.load(f))
#     return set()

def load_processed_files(processed_files_path):
    if os.path.exists(processed_files_path):
        try:
            with open(processed_files_path, 'r') as f:
                data = f.read().strip()
                if data:  # Check if the file has content
                    return set(json.loads(data))
        except json.JSONDecodeError:
            print(f"Error: {processed_files_path} contains invalid JSON. Starting with an empty set.")
    return set()


def process_folders_parallel(base_directory_path, output_base_path, timings_file, processed_files, num_workers=20):
    directories = [os.path.join(base_directory_path, d) for d in os.listdir(base_directory_path) if os.path.isdir(os.path.join(base_directory_path, d))]
    with Pool(num_workers) as pool:
        results = [pool.apply_async(process_directory, (directory, output_base_path, timings_file)) for directory in directories]
        folder_timings = [result.get() for result in results]

    # Save folder timings to the timings file
    with open(timings_file, 'a') as f:
        f.write("Overall Folder Timings:\n")
        for timing in folder_timings:
            for folder, elapsed_time in timing.items():
                f.write(f"{folder}: {elapsed_time:.2f} seconds\n")
if __name__ == "__main__":
    current_day = datetime.now().strftime("%d-%m-%y")
    # previous_day = (datetime.now() - timedelta(days=1)).strftime("%d-%m-%y")
    # Use f-strings to correctly format the paths
    base_directory_path = f"/data/unzipdocument/dailydocument_{current_day}" 
    output_base_path = f"/data/txtfolder/dailydocument_{current_day}_txt"
    timings_file = f"/data/QAAPI/dailydocument_{current_day}_live"
    processed_files_path = '/data/tendergpt/processed_files.json'

    
    overall_start_time = time.time()
    processed_files = load_processed_files(processed_files_path)
    print(f"Processed files loaded: {len(processed_files)} files.")

    process_folders_parallel(base_directory_path, output_base_path, timings_file, processed_files, num_workers=30)
    print("Folder processing completed.")

    save_processed_files(processed_files, processed_files_path)
    print("Processed files saved.")

    overall_elapsed_time = time.time() - overall_start_time
    with open(timings_file, 'a') as f:
        f.write(f"Overall elapsed time for this cycle: {overall_elapsed_time:.2f} seconds.\n")

    print("Overall process completed.")


